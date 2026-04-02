#!/usr/bin/env python3
"""Build Virgil_personality_synthesis.docx from the Markdown workbook."""

from __future__ import annotations

import sys
from pathlib import Path


def main() -> int:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        sys.stderr.write(
            "Missing python-docx. Install: pip install -r scripts/requirements-personality-docx.txt\n"
        )
        return 1

    root = Path(__file__).resolve().parent.parent
    md_path = root / "docs" / "personality" / "Virgil_personality_synthesis.md"
    out_path = root / "docs" / "personality" / "Virgil_personality_synthesis.docx"

    if not md_path.is_file():
        sys.stderr.write(f"Source not found: {md_path}\n")
        return 1

    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    normal = doc.styles["Normal"]
    if normal.font is not None:
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

    i = 0
    in_code = False
    code_buf: list[str] = []

    def is_table_row(s: str) -> bool:
        t = s.strip()
        return t.startswith("|") and t.endswith("|") and t.count("|") >= 2

    def is_table_separator(s: str) -> bool:
        t = s.strip().strip("|")
        return bool(t) and all(part.strip().replace("-", "").replace(":", "") == "" for part in t.split("|"))

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if is_table_row(line):
            rows: list[list[str]] = []
            while i < len(lines) and is_table_row(lines[i]):
                row_line = lines[i].strip()
                if is_table_separator(row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split("|")[1:-1]]
                rows.append(cells)
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                for r in rows:
                    while len(r) < ncol:
                        r.append("")
                table = doc.add_table(rows=len(rows), cols=ncol)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(ncol):
                        table.rows[ri].cells[ci].text = row[ci] if ci < len(row) else ""
            continue

        if line.strip() == "":
            i += 1
            continue

        buf: list[str] = []
        while i < len(lines):
            ln = lines[i]
            if ln.strip() == "":
                break
            if ln.strip().startswith("```"):
                break
            if ln.startswith("### ") or ln.startswith("## ") or ln.startswith("# "):
                break
            if ln.strip() == "---":
                break
            if is_table_row(ln):
                break
            buf.append(ln)
            i += 1
        if buf:
            doc.add_paragraph("\n".join(buf))
        continue

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    sys.stdout.write(f"Wrote {out_path}\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
